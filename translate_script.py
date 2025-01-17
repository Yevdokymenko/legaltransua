import logging
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
import openai
from transformers import MarianMTModel, MarianTokenizer
import time
from datetime import datetime
from googletrans import Translator, LANGUAGES
import os
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
import re
from deep_translator import GoogleTranslator
from dotenv import load_dotenv

# Завантаження змінних середовища з файлу .env
load_dotenv(dotenv_path="key.env")

# Призначення ключа OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")

# Налаштування логування
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def extract_text_from_docx(file_path):
    """Витягує текст із DOCX-файлу."""
    doc = docx.Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def extract_text_from_pdf(file_path):
    """Екстрагує текст із PDF-файлу."""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return [line.strip() for line in text.splitlines() if line.strip()]


def extract_text_from_html(url):
    """Екстрагує текст із веб-сторінки."""
    response = requests.get(url)
    if response.status_code != 200:
        raise Exception(f"Не вдалося завантажити сторінку: {url}")
    
    soup = BeautifulSoup(response.content, "html.parser")
    paragraphs = soup.find_all("p")
    return [para.get_text().strip() for para in paragraphs if para.get_text().strip()]

def extract_text(source):
    """Визначає тип джерела і екстрагує текст."""
    if isinstance(source, str):  # Перевірка на тип рядка (шлях до файлу або URL)
        if source.startswith("http"):
            logging.info("Джерело визначено як веб-сторінка.")
            return extract_text_from_html(source)
        elif source.endswith(".pdf"):
            logging.info("Джерело визначено як PDF-файл.")
            return extract_text_from_pdf(source)
        elif source.endswith(".docx"):
            logging.info("Джерело визначено як DOCX-файл.")
            return extract_text_from_docx(source)
    raise ValueError("Формат файлу не підтримується. Підтримуються DOCX, PDF або URL.")

def choose_directory():
    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    return output_dir

def sanitize_filename(filename):
    """Очищає ім'я файлу від недопустимих символів."""
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def extract_text_from_url(url):
    """Функція для витягнення тексту з веб-сторінки."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        # Отримати всі абзаци тексту
        paragraphs = [p.text.strip() for p in soup.find_all('p') if p.text.strip()]
        return paragraphs
    except requests.exceptions.RequestException as e:
        return f"Помилка при завантаженні URL: {e}"

def split_text_into_chunks(text, max_length=500):
    """Розбиває текст на частини для перекладу."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 > max_length:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
        else:
            current_chunk.append(word)
            current_length += len(word) + 1

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

def translate_text_google(text, max_retries=3):
    for attempt in range(max_retries):
        try:
            translated = GoogleTranslator(source='en', target='uk').translate(text)
            return translated
        except Exception as e:
            logging.warning(
                f"Google Translator Error (attempt {attempt + 1}/{max_retries}): {e}"
            )
            time.sleep(2 ** attempt)
    logging.error("Google Translate: Помилка після кількох спроб")
    return "Помилка перекладу через Google Translator"

def translate_text_marian(text, tokenizer, model):
    """Перекладає текст через MarianMT."""
    try:
        inputs = tokenizer([text], return_tensors="pt", padding=True, truncation=True)
        translated = model.generate(**inputs)
        return tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
    except Exception as e:
        logging.warning(f"MarianMT Error: {e}")
        return "Помилка перекладу"

def translate_text_openai(text, max_retries=3):
    """Перекладає текст через OpenAI GPT-3.5 Turbo з повторними спробами."""
    for attempt in range(max_retries):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Translate the following text to Ukrainian."},
                    {"role": "user", "content": text},
                ],
            )
            return response.choices[0].message["content"].strip()
        except Exception as e:
            logging.warning(f"OpenAI API Error (attempt {attempt + 1}/{max_retries}): {e}")
            time.sleep(2 ** attempt + 1)  # Експоненційний відкат
    return "Помилка перекладу"

def set_table_border(table):
    """Встановлює межі таблиці."""
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 4pt border
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def setup_document_orientation(doc):
    """Налаштовує горизонтальну орієнтацію документа, вузькі поля та номери сторінок."""
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Додаємо номери сторінок
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Додаємо поле для номера сторінки
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), "PAGE")
    run = paragraph.add_run()
    run._r.append(fld_simple)

    # Додаємо текст " of "
    run = paragraph.add_run(" of ")

    # Додаємо поле для кількості сторінок
    fld_simple_total = OxmlElement('w:fldSimple')
    fld_simple_total.set(qn('w:instr'), "NUMPAGES")
    run._r.append(fld_simple_total)

def add_title(doc):
    """Додає заголовок із жирним текстом розміру 12 по центру."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Документ створено за допомогою скрипта перекладу LegalTransUA від BRDO")
    run.bold = True
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_translation_table(doc, paragraphs, google_translations, marian_translations, openai_translations):
    """Створює таблицю перекладів у DOCX-документі."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["№", "Оригінальний текст", "Google Translate", "MarianMT", "OpenAI GPT"]
    header_fill_color = "D9EAF7"  # Світло-блакитний
    row_number_fill_color = "E0E0E0"  # Світло-сірий

    # Додаємо заголовки колонок
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True

        # Заливка кольору
        cell._element.get_or_add_tcPr().append(create_shading_element(header_fill_color))

    # Заповнення таблиці
    for i, (para, g_trans, m_trans, o_trans) in enumerate(zip(paragraphs, google_translations, marian_translations, openai_translations)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = para
        row_cells[2].text = g_trans
        row_cells[3].text = m_trans
        row_cells[4].text = o_trans

        # Заливка для першої колонки
        row_cells[0]._element.get_or_add_tcPr().append(create_shading_element(row_number_fill_color))

        # Вирівнювання тексту по ширині
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Встановлення ширини колонок
    total_width = docx.shared.Inches(10)
    column_widths = [total_width * 0.04, total_width * 0.23, total_width * 0.23, total_width * 0.23, total_width * 0.23]
    for i, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = column_widths[i]

    # Закріплення заголовка таблиці на кожній сторінці
    tbl = table._element
    tblHeader = tbl.xpath(".//w:tr")[0]
    trPr = tblHeader.get_or_add_trPr()
    tblHeaderElement = OxmlElement("w:tblHeader")
    tblHeaderElement.set(qn("w:val"), "1")
    trPr.append(tblHeaderElement)

    return doc

def create_shading_element(color):
    """Створює елемент заливки комірки."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    return shading

def save_translation_document(source, paragraphs, google_translations, marian_translations, openai_translations):
    """Зберігає переклади в новий DOCX-документ."""
    doc = docx.Document()
    setup_document_orientation(doc)
    add_title(doc)
    doc.add_paragraph(f"Дата та час перекладу: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Додаємо таблицю
    create_translation_table(doc, paragraphs, google_translations, marian_translations, openai_translations)

    # Визначення назви файлу
    if source.startswith("http"):
        base_name = source.split("/")[-1]  # Беремо останню частину URL
    else:
        base_name = os.path.splitext(os.path.basename(source))[0]
    
    sanitized_name = sanitize_filename(base_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Вибір папки для збереження
    save_directory = choose_directory()
    output_file = os.path.join(
        save_directory, f"{sanitized_name} (Translated by LTU) {timestamp}.docx"
    )

    # Збереження файлу
    doc.save(output_file)
    logging.info(f"Документ збережено за адресою: {output_file}")


def process_document(source, tokenizer, model):
    """Основна функція для обробки документа."""
    paragraphs = extract_text(source)
    logging.info(f"Знайдено абзаців: {len(paragraphs)}")

    # Ініціалізація MarianMT
    model_name = "Helsinki-NLP/opus-mt-en-uk"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)

    google_translations = [""] * len(paragraphs)
    marian_translations = [""] * len(paragraphs)
    openai_translations = [""] * len(paragraphs)

    with ThreadPoolExecutor(max_workers=10) as executor:
        google_futures = {executor.submit(translate_text_google, para): idx for idx, para in enumerate(paragraphs)}
        for future in tqdm(as_completed(google_futures), total=len(google_futures), desc="Google Translate", unit="paragraph"):
            idx = google_futures[future]
            google_translations[idx] = future.result()

        marian_futures = {executor.submit(translate_text_marian, para, tokenizer, model): idx for idx, para in enumerate(paragraphs)}
        for future in tqdm(as_completed(marian_futures), total=len(marian_futures), desc="MarianMT", unit="paragraph"):
            idx = marian_futures[future]
            marian_translations[idx] = future.result()

        openai_futures = {executor.submit(translate_text_openai, para): idx for idx, para in enumerate(paragraphs)}
        for future in tqdm(as_completed(openai_futures), total=len(openai_futures), desc="OpenAI GPT", unit="paragraph"):
            idx = openai_futures[future]
            openai_translations[idx] = future.result()

    save_translation_document(source, paragraphs, google_translations, marian_translations, openai_translations)


if __name__ == "__main__":
    source = input("Введіть URL, шлях до PDF або DOCX-файлу: ").strip()
    process_document(source, tokenizer, model)